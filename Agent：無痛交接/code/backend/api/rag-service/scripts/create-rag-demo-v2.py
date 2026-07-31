from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[2] / "data" / "無痛交接Demo資料_v2"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
MUTED = "666666"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != 9360:
        raise ValueError("Table widths must total 9360 DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=11, color="000000", bold=False, italic=False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def create_document(title: str, subtitle: str, code: str, owner: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(header.add_run(f"HANDOFF REFERENCE  |  {code}"), size=9, color=MUTED, bold=True)
    add_page_number(section.footer.paragraphs[0])

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    set_run(kicker.add_run("INTERNAL OPERATIONS HANDOFF"), size=9, color=BLUE, bold=True)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(4)
    set_run(title_p.add_run(title), size=24, color="0B2545", bold=True)
    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(14)
    set_run(subtitle_p.add_run(subtitle), size=13, color=MUTED)

    metadata = add_table(doc, ["文件編號", "文件版本", "交接負責人", "基準日"], [[code, "v2.0", owner, "2026 年 8 月 1 日"]], [1800, 1500, 3060, 3000])
    metadata.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()
    return doc


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = text
        set_cell_shading(cell, HEADER_FILL)
        for run in cell.paragraphs[0].runs:
            set_run(run, size=10.5, color=DARK_BLUE, bold=True)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    repeat_header(table.rows[0])
    for values in rows:
        row = table.add_row()
        for index, text in enumerate(values):
            row.cells[index].text = text
            for paragraph in row.cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run(run, size=10.2)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_intro(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run(text), size=11, color="333333")


def build_nimbus() -> None:
    doc = create_document(
        "Project NIMBUS 平台維運交接",
        "部署、事故應變、備份與存取權限",
        "NIMBUS-HO-OPS-002",
        "陳昱豪／平台維運工程師",
    )
    add_intro(doc, "本文件提供 NIMBUS 平台的日常維運基準。任何事故處置均須保留工單、時間線與決策依據，不得在文件中記錄密碼或權杖。")
    doc.add_heading("1. 上線與部署狀態", level=1)
    add_table(doc, ["項目", "目前狀態", "操作規則"], [
        ["發布階段", "Canary 監控中，正式流量 25%", "連續觀察 24 小時後才可提升流量"],
        ["變更凍結窗", "每週五 18:00-22:00", "凍結期間只允許 P1 緊急修復"],
        ["回滾條件", "錯誤率連續 5 分鐘高於 2%", "由 Incident Commander 宣告並執行上一版映像回滾"],
        ["Staging 權限", "NIMBUS-Staging-Operator", "由平台 IAM 群組核准，個人帳號登入"],
    ], [2200, 2800, 4360])
    doc.add_heading("2. 事故應變", level=1)
    add_table(doc, ["情境", "時限／位置", "負責方式"], [
        ["P1 服務中斷", "10 分鐘內回應", "先建立 INC 工單，再通知當週 on-call"],
        ["事故協作室", "Slack #inc-nimbus", "所有指令與判斷同步記錄在頻道 thread"],
        ["狀態頁公告", "Incident Commander 核准後發布", "未核准前只更新內部事故時間線"],
        ["主要 on-call", "陳昱豪，分機 7312", "無回應 5 分鐘後改撥平台主管"],
    ], [2500, 3100, 3760])
    doc.add_heading("3. 備份與機密管理", level=1)
    add_table(doc, ["控制項", "設定", "驗證方式"], [
        ["資料庫備份", "每日 02:30，自動保留 14 天", "每週一抽查最新備份可還原"],
        ["Secrets 輪替", "Vault 路徑 sec/nimbus/prod", "每 90 天輪替；只保存路徑，不保存 secret 值"],
        ["稽核紀錄", "OPS/NIMBUS/Audit", "部署與權限異動紀錄至少保存 365 天"],
    ], [2500, 3300, 3560])
    doc.save(OUTPUT / "04_平台維運交接_Project_NIMBUS.docx")


def build_lantern() -> None:
    doc = create_document(
        "Project LANTERN 客戶成功交接",
        "續約、健康度、會議節奏與升級窗口",
        "LANTERN-HO-CS-004",
        "黃嘉文／資深客戶成功經理",
    )
    add_intro(doc, "LANTERN 是跨區企業客戶。對外回覆必須以客戶已確認的里程碑為準；內部風險判斷不可直接轉寄給客戶。")
    doc.add_heading("1. 帳戶健康與續約", level=1)
    add_table(doc, ["項目", "內容", "下一步"], [
        ["續約日", "2026 年 11 月 30 日", "提前 120 天啟動方案與預算確認"],
        ["目前健康度", "黃色；行動版 MAU 較上月下降 18%", "兩週內完成使用者訪談並提出採用計畫"],
        ["主要贊助人", "Evelyn Wang／營運副總", "重要範圍與預算變更須先取得她的書面確認"],
        ["流失警訊", "若 9 月 15 日前未通過 SSO pilot，列為高風險", "每週追蹤資安審查與測試缺陷"],
    ], [2200, 3600, 3560])
    doc.add_heading("2. 溝通節奏", level=1)
    add_table(doc, ["會議／渠道", "固定時間", "規則"], [
        ["每週進度同步", "每週三 10:30", "會前一天寄出阻塞項目與 owner"],
        ["季度業務回顧 QBR", "每季第一週週二 14:00", "由客戶成功經理彙整採用率、價值與風險"],
        ["緊急升級信箱", "cs-lantern-urgent@example.com", "只用於 P1、資料安全或合約即時風險"],
        ["支援語言", "繁體中文與英文", "正式摘要提供雙語版本"],
    ], [2500, 2900, 3960])
    doc.add_heading("3. 行動與環境", level=1)
    add_table(doc, ["項目", "負責人／代碼", "完成標準"], [
        ["NPS 負評追蹤", "黃嘉文", "48 小時內聯絡並記錄改善承諾"],
        ["客戶展示環境", "Tenant LTN-DEMO-07", "每次展示前清除測試個資並確認版本"],
        ["採用率報表", "CS/LANTERN/Adoption", "每週一更新並標註資料截止日"],
    ], [2700, 2900, 3760])
    doc.save(OUTPUT / "05_客戶成功交接_Project_LANTERN.docx")


def build_aurora() -> None:
    doc = create_document(
        "Project AURORA 採購財務交接",
        "供應商合約、請購、發票與稽核證據",
        "AURORA-HO-FIN-003",
        "林慧珊／採購與財務專員",
    )
    add_intro(doc, "本文件是 AURORA 採購與付款作業索引。所有核准都必須在正式系統留痕，聊天訊息不可取代請購或付款核准。")
    doc.add_heading("1. 合約與請購", level=1)
    add_table(doc, ["項目", "值", "控制規則"], [
        ["主要供應商", "NovaPay Systems", "供應商主檔異動須由採購與財務雙重覆核"],
        ["採購單", "AUR-PO-2026-117", "追加金額需建立變更單，不可覆寫原始 PO"],
        ["合約到期日", "2027 年 3 月 31 日", "到期前 120 天啟動續約或替代方案評估"],
        ["高額核准門檻", "單筆超過新台幣 150,000 元須由 CFO 核准", "拆單規避門檻視為控制缺失"],
    ], [2500, 2900, 3960])
    doc.add_heading("2. 發票與付款", level=1)
    add_table(doc, ["項目", "作業方式", "例外處理"], [
        ["付款條件", "NET45；驗收完成日開始計算", "未完成驗收不得提前起算"],
        ["發票存放", "FIN/AP/AURORA/2026", "檔名使用供應商_發票號碼_日期"],
        ["發票不符", "暫停入帳，1 個工作天內通知財務", "保留原始發票與差異說明，不得自行修改"],
        ["作業負責人", "林慧珊", "請假時由財務代理人依同一控制流程處理"],
    ], [2200, 3700, 3460])
    doc.add_heading("3. 授權與稽核", level=1)
    add_table(doc, ["控制項", "規則", "保存位置／期限"], [
        ["軟體授權", "共 48 席；帳號閒置超過 30 天即回收", "ITAM/AURORA/Licenses"],
        ["付款證據", "PO、驗收、發票與核准紀錄必須可相互勾稽", "稽核證據保存 7 年"],
        ["銀行資料變更", "必須回撥供應商既有聯絡人確認", "FIN/Vendor-Control/Callbacks"],
    ], [2500, 3960, 2900])
    doc.save(OUTPUT / "06_採購財務交接_Project_AURORA.docx")


if __name__ == "__main__":
    OUTPUT.mkdir(parents=True, exist_ok=True)
    build_nimbus()
    build_lantern()
    build_aurora()
    print(f"Created {len(list(OUTPUT.glob('0[4-6]_*.docx')))} DOCX fixtures.")
